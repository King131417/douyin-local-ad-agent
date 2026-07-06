"""
Web Dashboard v3 — 双驾驶舱
  驾驶舱一：素材投资决策驾驶舱 (Material Cockpit)
    - 决策总览 / 账户穿透 / 素材全榜 / 优化建议
  驾驶舱二：投流账户驾驶舱 (Account Cockpit)
    - 全局KPI / 账户趋势 / 项目分析 / 单元排行
"""

import json
import logging
import os
import sqlite3
from datetime import date, timedelta

from flask import Flask, jsonify, render_template, request, send_file, send_from_directory

from src.analysis.kpi import KPIAnalyzer
from src.analysis.material_decision import MaterialDecisionEngine
from src.analysis.quality_radar import QualityRadar
from src.pipeline.storage import Storage
from src.pipeline.scheduler import AdDataScheduler, sync_status

logger = logging.getLogger(__name__)

app = Flask(__name__,
            template_folder=os.path.join(os.path.dirname(__file__), 'templates'))
kpi = KPIAnalyzer()
engine = MaterialDecisionEngine()
storage = Storage()
radar = QualityRadar()

# ── Helpers ─────────────────────────────────────────────────────
_scheduler: AdDataScheduler | None = None

def _start_scheduler():
    global _scheduler
    _scheduler = AdDataScheduler()
    _scheduler.start()
    logger.info("Auto-sync scheduler started (daily 9:30 AM Beijing time)")
    _scheduler.trigger_now()


def _default_date() -> str:
    return storage.get_latest_date("material_reports") or (
        date.today() - timedelta(days=1)
    ).strftime("%Y-%m-%d")


def _db_path() -> str:
    return os.path.join(
        os.path.dirname(os.path.dirname(os.path.dirname(__file__))),
        "data", "ad_data.db"
    )


def _get_account_name_map() -> dict:
    """Load account_id -> name from accounts table."""
    conn = sqlite3.connect(_db_path())
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    c.execute("SELECT account_id, name FROM accounts")
    m = {r["account_id"]: r["name"] for r in c.fetchall()}
    conn.close()
    return m


# ── Material Cockpit Routes ──────────────────────────────────────

@app.route("/")
def dashboard():
    return render_template('dashboard.html')


@app.route("/static/<path:filename>")
def static_files(filename):
    static_dir = os.path.join(os.path.dirname(__file__), "static")
    return send_from_directory(static_dir, filename)


@app.route("/api/summary")
def api_summary():
    days = int(request.args.get("days", 7))
    trend = kpi.trend(days)
    today = trend[-1] if trend else {}
    return jsonify({"today": today, "trend": trend})


@app.route("/api/decision")
def api_decision():
    start_date = request.args.get("start_date")
    end_date = request.args.get("end_date")
    date_str = request.args.get("date") or _default_date()
    result = engine.analyze(date_str=date_str, start_date=start_date, end_date=end_date)
    return jsonify(result)


@app.route("/api/account/<account_id>")
def api_account_detail(account_id: str):
    start_date = request.args.get("start_date")
    end_date = request.args.get("end_date")
    date_str = request.args.get("date") or _default_date()
    result = engine.get_account_detail(
        account_id, date_str=date_str,
        start_date=start_date, end_date=end_date,
    )
    return jsonify(result)


@app.route("/api/ranking")
def api_ranking():
    start_date = request.args.get("start_date")
    end_date = request.args.get("end_date")
    date_str = request.args.get("date") or _default_date()
    ranking = engine.get_full_ranking(date_str=date_str, start_date=start_date, end_date=end_date)
    return jsonify({"materials": ranking, "total": len(ranking)})


@app.route("/api/material/<material_id>")
def api_material_detail(material_id: str):
    start_date = request.args.get("start_date")
    end_date = request.args.get("end_date")
    conn = sqlite3.connect(_db_path())
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    if start_date and end_date:
        c.execute(
            "SELECT * FROM material_reports WHERE material_id = ? AND stat_date BETWEEN ? AND ? ORDER BY stat_date ASC",
            (material_id, start_date, end_date),
        )
    else:
        c.execute(
            "SELECT * FROM material_reports WHERE material_id = ? ORDER BY stat_date ASC",
            (material_id,),
        )
    rows = c.fetchall()
    conn.close()
    if not rows:
        return jsonify({"error": "material not found"}), 404

    # Build daily trend
    daily = []
    totals = {"cost": 0, "show": 0, "click": 0, "clue": 0, "consult": 0, "convert": 0}
    for r in rows:
        daily.append({
            "date": r["stat_date"],
            "cost": round(r["stat_cost"] or 0, 2),
            "show": r["show_cnt"] or 0,
            "click": r["click_cnt"] or 0,
            "ctr": round(r["ctr"] or 0, 1),
            "clue": r["clue_message_count"] or 0,
            "consult": r["message_action_cnt"] or 0,
            "convert": r["convert_cnt"] or 0,
        })
        totals["cost"] += r["stat_cost"] or 0
        totals["show"] += r["show_cnt"] or 0
        totals["click"] += r["click_cnt"] or 0
        totals["clue"] += r["clue_message_count"] or 0
        totals["consult"] += r["message_action_cnt"] or 0
        totals["convert"] += r["convert_cnt"] or 0

    # Collect attribution info safely (sqlite3.Row has no .get())
    promos = set()
    projects = set()
    first_with_attribution = None
    for r in rows:
        try:
            pid = r["promotion_id"]
            pname = r["promotion_name"]
            if pid and pname:
                promos.add(str(pname))
                if first_with_attribution is None:
                    first_with_attribution = r
        except (IndexError, KeyError):
            pass
        try:
            pjid = r["project_id"]
            pjname = r["project_name"]
            if pjid and pjname:
                projects.add(str(pjname))
        except (IndexError, KeyError):
            pass

    # First row for basic info: prefer one with attribution, fall back to first
    first = first_with_attribution or rows[0]

    # Helper for safe column access
    def _safe(row, col):
        try:
            v = row[col]
            return str(v) if v else ""
        except (IndexError, KeyError):
            return ""

    # Build response dict explicitly (avoid nested comprehension issues)
    result = {}
    result["material_id"] = str(first["material_id"])
    result["material_name"] = str(first["material_name"])
    result["material_type"] = _safe(first, "material_type")
    result["account_id"] = str(first["account_id"])
    result["promotion_id"] = _safe(first, "promotion_id")
    result["promotion_name"] = _safe(first, "promotion_name")
    result["project_id"] = _safe(first, "project_id")
    result["project_name"] = _safe(first, "project_name")
    result["all_promotions"] = sorted(promos)
    result["all_projects"] = sorted(projects)
    result["daily"] = daily
    result["totals"] = totals

    return jsonify(result)


@app.route("/api/material/<material_id>/creative")
def api_material_creative(material_id: str):
    conn = sqlite3.connect(_db_path())
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    c.execute(
        "SELECT DISTINCT account_id FROM material_reports WHERE material_id = ? LIMIT 1",
        (material_id,),
    )
    row = c.fetchone()
    conn.close()
    if not row:
        return jsonify({"error": "material not found"}), 404
    account_id = row["account_id"]
    try:
        from src.api.client import OceanEngineClient
        client = OceanEngineClient()
        today_str = _default_date()
        from datetime import date as _date, timedelta as _td
        start = (_date.today() - _td(days=7)).strftime("%Y-%m-%d")
        promo_result = client._get("/local/report/promotion/get/", {
            "local_account_id": int(account_id),
            "start_date": start,
            "end_date": today_str,
            "page": 1, "page_size": 50,
            "metrics": ["stat_cost"],
        })
        if promo_result.get("code") != 0:
            return jsonify({"error": "promotion API failed"}), 502
        promotions = promo_result.get("data", {}).get("promotion_list", [])
        unique_pids = list(set(p.get("promotion_id") for p in promotions))
        for pid in unique_pids:
            detail = client._get("/local/promotion/detail/", {
                "local_account_id": int(account_id),
                "promotion_id": pid,
            })
            if detail.get("code") != 0:
                continue
            pm = detail.get("data", {}).get("procedural_material", {})
            for key in ("video_material_list", "image_material_list", "title_material_list"):
                for m in pm.get(key, []):
                    if str(m.get("lego_material_id", "")) == str(material_id):
                        return jsonify({
                            "creative_material_id": str(m.get("material_id", "")),
                            "video_id": m.get("video_id"),
                            "image_id": m.get("image_id"),
                            "material_type": key.split("_")[0],
                            "promotion_id": pid,
                        })
        return jsonify({"error": "creative mapping not found"}), 404
    except Exception as e:
        logger.warning("Creative lookup failed for %s: %s", material_id, e)
        return jsonify({"error": "lookup failed"}), 502


@app.route("/api/projects/<account_id>")
def api_projects(account_id: str):
    try:
        from src.api.client import OceanEngineClient
        client = OceanEngineClient()
        promotions = client.get_promotion_report(
            account_id, _default_date(), _default_date(),
        )
        seen = set()
        projects = []
        for p in promotions:
            cid = p.get("project_id", "")
            if cid and cid not in seen:
                seen.add(cid)
                projects.append({
                    "project_id": cid,
                    "project_name": p.get("project_name", ""),
                    "promotion_id": p.get("promotion_id", ""),
                })
        return jsonify({"projects": projects})
    except Exception as e:
        logger.warning("Failed to get projects for %s: %s", account_id, e)
        return jsonify({"projects": []})


@app.route("/api/sync/status")
def api_sync_status():
    return jsonify(sync_status.to_dict())


@app.route("/api/sync/trigger", methods=["POST"])
def api_sync_trigger():
    if sync_status.is_syncing:
        return jsonify({"ok": False, "message": "同步正在进行中，请稍后…"}), 409
    if _scheduler is None:
        return jsonify({"ok": False, "message": "调度器未初始化"}), 500
    _scheduler.trigger_now()
    return jsonify({"ok": True, "message": "同步已触发，稍后刷新数据"})


@app.route("/api/sync/backfill", methods=["POST"])
def api_sync_backfill():
    if sync_status.is_syncing:
        return jsonify({"ok": False, "message": "同步正在进行中，请稍后…"}), 409
    if _scheduler is None:
        return jsonify({"ok": False, "message": "调度器未初始化"}), 500
    days = int(request.args.get("days", 30))
    _scheduler.trigger_backfill(days=days)
    return jsonify({"ok": True, "message": f"已触发 {days} 天历史数据回填，正在后台执行…"})


# ── Account Cockpit Routes ───────────────────────────────────────

@app.route("/api/ad/overview")
def api_ad_overview():
    """
    Account cockpit: global KPIs + per-account summary.
    ?start_date=&end_date=  (default: latest 7 days)
    """
    end = request.args.get("end_date") or _default_date()
    start = request.args.get("start_date") or (
        date.fromisoformat(end) - timedelta(days=6)
    ).strftime("%Y-%m-%d")

    name_map = _get_account_name_map()
    conn = sqlite3.connect(_db_path())
    conn.row_factory = sqlite3.Row
    c = conn.cursor()

    # Global KPIs
    c.execute("""
        SELECT
            SUM(stat_cost)           AS cost,
            SUM(show_cnt)            AS show,
            SUM(click_cnt)           AS click,
            SUM(clue_message_count)  AS clue,
            SUM(message_action_cnt)  AS consult,
            SUM(convert_cnt)         AS convert,
            AVG(ctr)                 AS avg_ctr,
            AVG(cpm)                 AS avg_cpm,
            COUNT(DISTINCT account_id) AS active_accounts
        FROM account_reports
        WHERE stat_date BETWEEN ? AND ?
    """, (start, end))
    g = dict(conn.execute("""
        SELECT
            SUM(stat_cost) AS cost,
            SUM(show_cnt) AS show,
            SUM(click_cnt) AS click,
            SUM(clue_message_count) AS clue,
            SUM(message_action_cnt) AS consult,
            SUM(convert_cnt) AS convert,
            COUNT(DISTINCT account_id) AS active_accounts,
            COUNT(DISTINCT stat_date) AS days
        FROM account_reports
        WHERE stat_date BETWEEN ? AND ?
          AND delivery_type = 'total'
    """, (start, end)).fetchone())

    total_cost = g["cost"] or 0
    total_clue = g["clue"] or 0
    total_consult = g["consult"] or 0
    g["cpa_clue"] = round(total_cost / total_clue, 2) if total_clue else 0
    g["cpa_consult"] = round(total_cost / total_consult, 2) if total_consult else 0
    g["ctr"] = round(
        (g["click"] or 0) / (g["show"] or 1) * 100, 2
    )

    # Per-account summary
    rows = c.execute("""
        SELECT
            account_id,
            SUM(stat_cost)           AS cost,
            SUM(show_cnt)            AS show,
            SUM(click_cnt)           AS click,
            SUM(clue_message_count)  AS clue,
            SUM(message_action_cnt)  AS consult,
            SUM(convert_cnt)         AS convert,
            AVG(cpm)                 AS avg_cpm,
            COUNT(DISTINCT stat_date) AS active_days
        FROM account_reports
        WHERE stat_date BETWEEN ? AND ?
          AND delivery_type = 'total'
        GROUP BY account_id
        ORDER BY cost DESC
    """, (start, end)).fetchall()

    accounts = []
    for r in rows:
        cost = r["cost"] or 0
        clue = r["clue"] or 0
        consult = r["consult"] or 0
        ctr = round((r["click"] or 0) / (r["show"] or 1) * 100, 2)
        accounts.append({
            "account_id": r["account_id"],
            "account_name": name_map.get(r["account_id"], r["account_id"][-8:]),
            "cost": round(cost, 2),
            "show": r["show"] or 0,
            "click": r["click"] or 0,
            "clue": clue,
            "consult": consult,
            "convert": r["convert"] or 0,
            "ctr": ctr,
            "avg_cpm": round(r["avg_cpm"] or 0, 2),
            "cpa_clue": round(cost / clue, 2) if clue else 0,
            "cpa_consult": round(cost / consult, 2) if consult else 0,
            "active_days": r["active_days"] or 0,
            "cost_share": round(cost / total_cost * 100, 1) if total_cost else 0,
        })

    conn.close()
    return jsonify({
        "start": start, "end": end,
        "global": g,
        "accounts": accounts,
    })


@app.route("/api/ad/trend")
def api_ad_trend():
    """
    Daily trend for all accounts (or a single account).
    ?start_date=&end_date=&account_id=
    """
    end = request.args.get("end_date") or _default_date()
    start = request.args.get("start_date") or (
        date.fromisoformat(end) - timedelta(days=13)
    ).strftime("%Y-%m-%d")
    account_id = request.args.get("account_id", "")

    name_map = _get_account_name_map()
    conn = sqlite3.connect(_db_path())
    conn.row_factory = sqlite3.Row

    if account_id:
        rows = conn.execute("""
            SELECT stat_date, account_id,
                   stat_cost AS cost, show_cnt AS show, click_cnt AS click,
                   clue_message_count AS clue, message_action_cnt AS consult,
                   convert_cnt AS convert, ctr, cpm
            FROM account_reports
            WHERE stat_date BETWEEN ? AND ? AND account_id = ?
              AND delivery_type = 'total'
            ORDER BY stat_date
        """, (start, end, account_id)).fetchall()
    else:
        rows = conn.execute("""
            SELECT stat_date,
                   SUM(stat_cost)           AS cost,
                   SUM(show_cnt)            AS show,
                   SUM(click_cnt)           AS click,
                   SUM(clue_message_count)  AS clue,
                   SUM(message_action_cnt)  AS consult,
                   SUM(convert_cnt)         AS convert,
                   AVG(ctr)                 AS ctr,
                   AVG(cpm)                 AS cpm
            FROM account_reports
            WHERE stat_date BETWEEN ? AND ?
              AND delivery_type = 'total'
            GROUP BY stat_date
            ORDER BY stat_date
        """, (start, end)).fetchall()

    # Also get per-account daily (for stacked chart)
    acc_rows = conn.execute("""
        SELECT stat_date, account_id,
               SUM(stat_cost) AS cost,
               SUM(clue_message_count) AS clue
        FROM account_reports
        WHERE stat_date BETWEEN ? AND ?
          AND delivery_type = 'total'
        GROUP BY stat_date, account_id
        ORDER BY stat_date, cost DESC
    """, (start, end)).fetchall()
    conn.close()

    trend = [
        {
            "date": r["stat_date"],
            "cost": round(r["cost"] or 0, 2),
            "show": r["show"] or 0,
            "click": r["click"] or 0,
            "clue": r["clue"] or 0,
            "consult": r["consult"] or 0,
            "convert": r["convert"] or 0,
            "ctr": round(r["ctr"] or 0, 2),
            "cpm": round(r["cpm"] or 0, 2),
        }
        for r in rows
    ]

    # Build stacked series: {account_id: [{date, cost}, ...]}
    stacked: dict = {}
    for r in acc_rows:
        aid = r["account_id"]
        if aid not in stacked:
            stacked[aid] = {"name": name_map.get(aid, aid[-8:]), "data": []}
        stacked[aid]["data"].append({
            "date": r["stat_date"],
            "cost": round(r["cost"] or 0, 2),
            "clue": r["clue"] or 0,
        })

    return jsonify({"trend": trend, "stacked": list(stacked.values())})


@app.route("/api/ad/projects")
def api_ad_projects():
    """
    Project-level analysis from promotion_reports.
    ?start_date=&end_date=&account_id=
    """
    end = request.args.get("end_date") or _default_date()
    start = request.args.get("start_date") or (
        date.fromisoformat(end) - timedelta(days=6)
    ).strftime("%Y-%m-%d")
    account_id = request.args.get("account_id", "")

    name_map = _get_account_name_map()
    conn = sqlite3.connect(_db_path())
    conn.row_factory = sqlite3.Row

    where = "WHERE stat_date BETWEEN ? AND ?"
    params: list = [start, end]
    if account_id:
        where += " AND account_id = ?"
        params.append(account_id)

    rows = conn.execute(f"""
        SELECT
            project_id,
            project_name,
            account_id,
            SUM(stat_cost)           AS cost,
            SUM(show_cnt)            AS show,
            SUM(click_cnt)           AS click,
            SUM(clue_message_count)  AS clue,
            SUM(message_action_cnt)  AS consult,
            SUM(convert_cnt)         AS convert,
            COUNT(DISTINCT promotion_id) AS promo_count,
            COUNT(DISTINCT stat_date)    AS active_days
        FROM promotion_reports
        {where}
        GROUP BY project_id, account_id
        ORDER BY cost DESC
    """, params).fetchall()
    conn.close()

    projects = []
    for r in rows:
        cost = r["cost"] or 0
        clue = r["clue"] or 0
        consult = r["consult"] or 0
        ctr = round((r["click"] or 0) / (r["show"] or 1) * 100, 2)
        projects.append({
            "project_id": r["project_id"],
            "project_name": r["project_name"] or r["project_id"],
            "account_id": r["account_id"],
            "account_name": name_map.get(r["account_id"], r["account_id"][-8:]),
            "cost": round(cost, 2),
            "show": r["show"] or 0,
            "click": r["click"] or 0,
            "clue": clue,
            "consult": consult,
            "convert": r["convert"] or 0,
            "ctr": ctr,
            "cpa_clue": round(cost / clue, 2) if clue else 0,
            "cpa_consult": round(cost / consult, 2) if consult else 0,
            "promo_count": r["promo_count"] or 0,
            "active_days": r["active_days"] or 0,
        })

    return jsonify({"projects": projects, "start": start, "end": end})


@app.route("/api/ad/promotions")
def api_ad_promotions():
    """
    Promotion-unit level ranking.
    ?start_date=&end_date=&account_id=&project_id=&limit=50
    """
    end = request.args.get("end_date") or _default_date()
    start = request.args.get("start_date") or (
        date.fromisoformat(end) - timedelta(days=6)
    ).strftime("%Y-%m-%d")
    account_id = request.args.get("account_id", "")
    project_id = request.args.get("project_id", "")
    limit = int(request.args.get("limit", 50))

    name_map = _get_account_name_map()
    conn = sqlite3.connect(_db_path())
    conn.row_factory = sqlite3.Row

    where = "WHERE stat_date BETWEEN ? AND ?"
    params: list = [start, end]
    if account_id:
        where += " AND account_id = ?"
        params.append(account_id)
    if project_id:
        where += " AND project_id = ?"
        params.append(project_id)

    rows = conn.execute(f"""
        SELECT
            promotion_id,
            promotion_name,
            promotion_status,
            project_id,
            project_name,
            account_id,
            SUM(stat_cost)           AS cost,
            SUM(show_cnt)            AS show,
            SUM(click_cnt)           AS click,
            SUM(clue_message_count)  AS clue,
            SUM(message_action_cnt)  AS consult,
            SUM(convert_cnt)         AS convert,
            COUNT(DISTINCT stat_date) AS active_days
        FROM promotion_reports
        {where}
        GROUP BY promotion_id
        ORDER BY cost DESC
        LIMIT ?
    """, params + [limit]).fetchall()
    conn.close()

    promos = []
    for r in rows:
        cost = r["cost"] or 0
        clue = r["clue"] or 0
        consult = r["consult"] or 0
        ctr = round((r["click"] or 0) / (r["show"] or 1) * 100, 2)
        promos.append({
            "promotion_id": r["promotion_id"],
            "promotion_name": r["promotion_name"] or r["promotion_id"],
            "promotion_status": r["promotion_status"] or "",
            "project_id": r["project_id"],
            "project_name": r["project_name"] or "",
            "account_id": r["account_id"],
            "account_name": name_map.get(r["account_id"], r["account_id"][-8:]),
            "cost": round(cost, 2),
            "show": r["show"] or 0,
            "click": r["click"] or 0,
            "clue": clue,
            "consult": consult,
            "convert": r["convert"] or 0,
            "ctr": ctr,
            "cpa_clue": round(cost / clue, 2) if clue else 0,
            "cpa_consult": round(cost / consult, 2) if consult else 0,
            "active_days": r["active_days"] or 0,
        })

    return jsonify({"promotions": promos, "start": start, "end": end})


# ── Account Report Routes (单账户综合报告) ─────────────────────

@app.route("/api/account_report/<account_id>")
def api_account_report(account_id: str):
    """
    单个账户的综合报告：KPI、每日趋势、项目排行、单元排行、素材排行。
    ?start_date=&end_date=
    """
    end = request.args.get("end_date") or _default_date()
    start = request.args.get("start_date") or (
        date.fromisoformat(end) - timedelta(days=6)
    ).strftime("%Y-%m-%d")

    name_map = _get_account_name_map()
    account_name = name_map.get(account_id, account_id[-8:])
    conn = sqlite3.connect(_db_path())
    conn.row_factory = sqlite3.Row
    c = conn.cursor()

    # ── Account KPIs ──
    acc_row = c.execute("""
        SELECT
            SUM(stat_cost)           AS cost,
            SUM(show_cnt)            AS show,
            SUM(click_cnt)           AS click,
            SUM(clue_message_count)  AS clue,
            SUM(message_action_cnt)  AS consult,
            SUM(convert_cnt)         AS convert,
            COUNT(DISTINCT stat_date) AS active_days
        FROM account_reports
        WHERE account_id = ? AND stat_date BETWEEN ? AND ?
          AND delivery_type = 'total'
    """, (account_id, start, end)).fetchone()

    cost = acc_row["cost"] or 0
    clue = acc_row["clue"] or 0
    consult = acc_row["consult"] or 0
    show_val = acc_row["show"] or 0
    click_val = acc_row["click"] or 0

    kpi = {
        "account_id": account_id,
        "account_name": account_name,
        "cost": round(cost, 2),
        "show": show_val,
        "click": click_val,
        "clue": clue,
        "consult": consult,
        "convert": acc_row["convert"] or 0,
        "ctr": round(click_val / (show_val or 1) * 100, 2),
        "cpm": round(cost / (show_val or 1) * 1000, 2),
        "lead_cost": round(cost / clue, 2) if clue else 0,
        "consult_cost": round(cost / consult, 2) if consult else 0,
        "lead_rate": round(clue / (consult or 1) * 100, 1),
        "active_days": acc_row["active_days"] or 0,
    }

    # ── Daily trend ──
    trend_rows = c.execute("""
        SELECT stat_date,
               stat_cost AS cost, show_cnt AS show, click_cnt AS click,
               clue_message_count AS clue, message_action_cnt AS consult,
               convert_cnt AS convert, ctr
        FROM account_reports
        WHERE account_id = ? AND stat_date BETWEEN ? AND ?
          AND delivery_type = 'total'
        ORDER BY stat_date
    """, (account_id, start, end)).fetchall()

    trend = [
        {
            "date": r["stat_date"],
            "cost": round(r["cost"] or 0, 2),
            "show": r["show"] or 0,
            "click": r["click"] or 0,
            "clue": r["clue"] or 0,
            "consult": r["consult"] or 0,
            "convert": r["convert"] or 0,
            "ctr": round(r["ctr"] or 0, 2),
        }
        for r in trend_rows
    ]

    # [DEPRECATED 2026-06-28] 通投/搜索拆分已废弃（按 SKILL.md 规范）
    # API 默认返回通投+搜索全量，不再做拆分展示。
    delivery_split = {}

    # ── Projects ──
    proj_rows = c.execute("""
        SELECT
            project_id, project_name,
            SUM(stat_cost)           AS cost,
            SUM(show_cnt)            AS show,
            SUM(click_cnt)           AS click,
            SUM(clue_message_count)  AS clue,
            SUM(message_action_cnt)  AS consult,
            SUM(convert_cnt)         AS convert,
            COUNT(DISTINCT promotion_id) AS promo_count,
            COUNT(DISTINCT stat_date)    AS active_days
        FROM promotion_reports
        WHERE account_id = ? AND stat_date BETWEEN ? AND ?
        GROUP BY project_id
        ORDER BY cost DESC
    """, (account_id, start, end)).fetchall()

    projects = []
    for r in proj_rows:
        pc = r["cost"] or 0
        pl = r["clue"] or 0
        projects.append({
            "project_id": r["project_id"],
            "project_name": r["project_name"] or r["project_id"],
            "cost": round(pc, 2),
            "show": r["show"] or 0,
            "click": r["click"] or 0,
            "clue": pl,
            "consult": r["consult"] or 0,
            "convert": r["convert"] or 0,
            "ctr": round((r["click"] or 0) / (r["show"] or 1) * 100, 2),
            "lead_cost": round(pc / pl, 2) if pl else 0,
            "promo_count": r["promo_count"] or 0,
            "active_days": r["active_days"] or 0,
        })

    # ── Promotions (top 30) ──
    promo_rows = c.execute("""
        SELECT
            promotion_id, promotion_name, promotion_status,
            project_id, project_name,
            SUM(stat_cost)           AS cost,
            SUM(show_cnt)            AS show,
            SUM(click_cnt)           AS click,
            SUM(clue_message_count)  AS clue,
            SUM(message_action_cnt)  AS consult,
            SUM(convert_cnt)         AS convert,
            COUNT(DISTINCT stat_date) AS active_days
        FROM promotion_reports
        WHERE account_id = ? AND stat_date BETWEEN ? AND ?
        GROUP BY promotion_id
        ORDER BY cost DESC
        LIMIT 30
    """, (account_id, start, end)).fetchall()

    promotions = []
    for r in promo_rows:
        prc = r["cost"] or 0
        prl = r["clue"] or 0
        promotions.append({
            "promotion_id": r["promotion_id"],
            "promotion_name": r["promotion_name"] or r["promotion_id"],
            "status": r["promotion_status"] or "",
            "project_name": r["project_name"] or r["project_id"],
            "cost": round(prc, 2),
            "show": r["show"] or 0,
            "click": r["click"] or 0,
            "clue": prl,
            "consult": r["consult"] or 0,
            "convert": r["convert"] or 0,
            "ctr": round((r["click"] or 0) / (r["show"] or 1) * 100, 2),
            "lead_cost": round(prc / prl, 2) if prl else 0,
            "active_days": r["active_days"] or 0,
        })

    # ── Materials (top 30) ──
    mat_rows = c.execute("""
        SELECT
            material_id, material_name, material_type,
            SUM(stat_cost)           AS cost,
            SUM(show_cnt)            AS show,
            SUM(click_cnt)           AS click,
            SUM(clue_message_count)  AS clue,
            SUM(message_action_cnt)  AS consult,
            SUM(convert_cnt)         AS convert,
            COUNT(DISTINCT stat_date) AS active_days,
            MAX(CASE WHEN promotion_id != '' THEN promotion_id END) as promotion_id,
            MAX(CASE WHEN promotion_id != '' THEN promotion_name END) as promotion_name,
            MAX(CASE WHEN project_id != '' THEN project_id END) as project_id,
            MAX(CASE WHEN project_id != '' THEN project_name END) as project_name
        FROM material_reports
        WHERE account_id = ? AND stat_date BETWEEN ? AND ?
        GROUP BY material_id
        ORDER BY cost DESC
        LIMIT 30
    """, (account_id, start, end)).fetchall()

    materials = []
    for r in mat_rows:
        mc = r["cost"] or 0
        ml = r["clue"] or 0
        mcons = r["consult"] or 0
        materials.append({
            "material_id": r["material_id"],
            "material_name": r["material_name"] or r["material_id"],
            "material_type": r["material_type"] or "",
            "cost": round(mc, 2),
            "show": r["show"] or 0,
            "click": r["click"] or 0,
            "clue": ml,
            "consult": mcons,
            "convert": r["convert"] or 0,
            "ctr": round((r["click"] or 0) / (r["show"] or 1) * 100, 2),
            "lead_cost": round(mc / ml, 2) if ml else 0,
            "lead_rate": round(ml / (mcons or 1) * 100, 1),
            "active_days": r["active_days"] or 0,
            # Attribution (v1.6)
            "promotion_id": r["promotion_id"] or "",
            "promotion_name": r["promotion_name"] or "",
            "project_id": r["project_id"] or "",
            "project_name": r["project_name"] or "",
        })

    conn.close()
    return jsonify({
        "start": start, "end": end,
        "kpi": kpi,
        "trend": trend,
        "delivery_split": delivery_split,
        "projects": projects,
        "promotions": promotions,
        "materials": materials,
    })


@app.route("/api/account_report/<account_id>/export")
def api_account_report_export(account_id: str):
    """
    导出单个账户报告为 Word 文档。
    ?start_date=&end_date=
    """
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from io import BytesIO

    end = request.args.get("end_date") or _default_date()
    start = request.args.get("start_date") or (
        date.fromisoformat(end) - timedelta(days=6)
    ).strftime("%Y-%m-%d")

    name_map = _get_account_name_map()
    account_name = name_map.get(account_id, account_id[-8:])
    conn = sqlite3.connect(_db_path())
    conn.row_factory = sqlite3.Row
    c = conn.cursor()

    # KPIs
    acc_row = c.execute("""
        SELECT SUM(stat_cost) AS cost, SUM(show_cnt) AS show, SUM(click_cnt) AS click,
               SUM(clue_message_count) AS clue, SUM(message_action_cnt) AS consult,
               SUM(convert_cnt) AS convert, COUNT(DISTINCT stat_date) AS active_days
        FROM account_reports
        WHERE account_id = ? AND stat_date BETWEEN ? AND ?
          AND delivery_type = 'total'
    """, (account_id, start, end)).fetchone()

    cost = acc_row["cost"] or 0
    clue = acc_row["clue"] or 0
    consult = acc_row["consult"] or 0
    show_val = acc_row["show"] or 0
    click_val = acc_row["click"] or 0
    convert_val = acc_row["convert"] or 0

    # [DEPRECATED 2026-06-28] Delivery type split removed (按 SKILL.md 规范)
    delivery_data = {}

    # Projects
    proj_rows = c.execute("""
        SELECT project_id, project_name,
               SUM(stat_cost) AS cost, SUM(clue_message_count) AS clue,
               SUM(message_action_cnt) AS consult, SUM(convert_cnt) AS convert,
               COUNT(DISTINCT promotion_id) AS promo_count
        FROM promotion_reports
        WHERE account_id = ? AND stat_date BETWEEN ? AND ?
        GROUP BY project_id ORDER BY cost DESC
    """, (account_id, start, end)).fetchall()

    # Promotions
    promo_rows = c.execute("""
        SELECT promotion_id, promotion_name, project_name,
               SUM(stat_cost) AS cost, SUM(clue_message_count) AS clue,
               SUM(convert_cnt) AS convert, SUM(show_cnt) AS show, SUM(click_cnt) AS click
        FROM promotion_reports
        WHERE account_id = ? AND stat_date BETWEEN ? AND ?
        GROUP BY promotion_id ORDER BY cost DESC LIMIT 30
    """, (account_id, start, end)).fetchall()

    # Materials
    mat_rows = c.execute("""
        SELECT material_id, material_name, material_type,
               SUM(stat_cost) AS cost, SUM(clue_message_count) AS clue,
               SUM(message_action_cnt) AS consult, SUM(convert_cnt) AS convert,
               SUM(show_cnt) AS show, SUM(click_cnt) AS click
        FROM material_reports
        WHERE account_id = ? AND stat_date BETWEEN ? AND ?
        GROUP BY material_id ORDER BY cost DESC LIMIT 30
    """, (account_id, start, end)).fetchall()

    conn.close()

    # ── Build Word Document ──
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    title = doc.add_heading(f'投流账户报告 — {account_name}', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f'数据周期：{start} ~ {end}    |    账户ID：{account_id}').font.size = Pt(9)

    doc.add_paragraph()

    # Section 1: KPI
    doc.add_heading('一、核心指标概览', level=1)

    kpi_data = [
        ("总消耗", f"¥{cost:,.0f}"),
        ("总展示", f"{show_val:,}"),
        ("总点击", f"{click_val:,}"),
        ("CTR", f"{click_val / (show_val or 1) * 100:.2f}%"),
        ("私信留资", f"{clue:,}"),
        ("私信咨询", f"{consult:,}"),
        ("留资成本", f"¥{cost / clue:.0f}" if clue else "--"),
        ("咨询成本", f"¥{cost / consult:.0f}" if consult else "--"),
        ("留咨率", f"{clue / (consult or 1) * 100:.1f}%"),
        ("活跃天数", f"{acc_row['active_days'] or 0}天"),
    ]

    kt = doc.add_table(rows=1, cols=4)
    kt.alignment = WD_TABLE_ALIGNMENT.CENTER
    kt.style = 'Light Grid Accent 1'
    for i, (label, value) in enumerate(kpi_data):
        row_idx = i // 2
        col_idx = (i % 2) * 2
        if row_idx >= len(kt.rows):
            kt.add_row()
        row = kt.rows[row_idx]
        row.cells[col_idx].text = label
        row.cells[col_idx + 1].text = value

    doc.add_paragraph()

    # [DEPRECATED 2026-06-28] 投放类型拆分章节已废弃（按 SKILL.md 规范）

    # Section 2: Projects
    doc.add_heading('二、项目排行', level=1)
    if proj_rows:
        pt = doc.add_table(rows=1, cols=6)
        pt.style = 'Light Grid Accent 1'
        pt.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(["项目名称", "消耗", "咨询", "留资", "投放单元数"]):
            pt.rows[0].cells[i].text = h
        for r in proj_rows:
            row = pt.add_row()
            row.cells[0].text = r["project_name"] or r["project_id"]
            row.cells[1].text = f'¥{r["cost"]:,.0f}' if r["cost"] else "0"
            row.cells[2].text = str(r["consult"] or 0)
            row.cells[3].text = str(r["clue"] or 0)
            row.cells[5].text = str(r["promo_count"] or 0)
    else:
        doc.add_paragraph("（无项目数据）")

    doc.add_paragraph()

    # Section 3: Promotions
    doc.add_heading('三、投放单元排行 (TOP 20)', level=1)
    if promo_rows:
        prt = doc.add_table(rows=1, cols=8)
        prt.style = 'Light Grid Accent 1'
        prt.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(["单元名称", "所属项目", "消耗", "展示", "点击", "留资", "留资成本"]):
            prt.rows[0].cells[i].text = h
        for r in promo_rows[:20]:
            row = prt.add_row()
            row.cells[0].text = (r["promotion_name"] or r["promotion_id"])[:30]
            row.cells[1].text = (r["project_name"] or "")[:20]
            row.cells[2].text = f'¥{r["cost"]:,.0f}' if r["cost"] else "0"
            row.cells[3].text = str(r["show"] or 0)
            row.cells[4].text = str(r["click"] or 0)
            prl_val = r["clue"] or 0
            row.cells[5].text = str(prl_val)
            row.cells[6].text = f'¥{r["cost"] / prl_val:.0f}' if prl_val and r["cost"] else "--"
    else:
        doc.add_paragraph("（无投放单元数据）")

    doc.add_paragraph()

    # Section 4: Materials
    doc.add_heading('四、素材排行 (TOP 20)', level=1)
    if mat_rows:
        mt = doc.add_table(rows=1, cols=8)
        mt.style = 'Light Grid Accent 1'
        mt.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(["素材名称", "类型", "消耗", "展示", "点击", "线索", "留咨率", "留资成本"]):
            mt.rows[0].cells[i].text = h
        for r in mat_rows[:20]:
            row = mt.add_row()
            row.cells[0].text = (r["material_name"] or r["material_id"])[:30]
            row.cells[1].text = r["material_type"] or ""
            row.cells[2].text = f'¥{r["cost"]:,.0f}' if r["cost"] else "0"
            row.cells[3].text = str(r["show"] or 0)
            row.cells[4].text = str(r["click"] or 0)
            mv_cost = r["cost"] or 0
            mv_clue = r["clue"] or 0
            mv_consult = r["consult"] or 0
            row.cells[5].text = str(mv_clue)
            row.cells[6].text = f'{mv_clue / (mv_consult or 1) * 100:.1f}%'
            row.cells[7].text = f'¥{mv_cost / mv_clue:.0f}' if mv_clue else "--"

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run(f'报告生成时间：{date.today().strftime("%Y-%m-%d")}').font.size = Pt(8)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f'{account_name}_投流报告_{start}_{end}.docx'
    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename,
    )


# ── Radar / Quality Diagnosis Routes ─────────────────────────────

@app.route("/api/radar/accounts")
def api_radar_accounts():
    """投手诊断：账户级质量雷达"""
    days = int(request.args.get("days", 7))
    start_date = request.args.get("start_date")
    end_date = request.args.get("end_date")
    return jsonify(radar.accounts(days, start_date=start_date, end_date=end_date))


@app.route("/api/radar/promotions/<account_id>")
def api_radar_promotions(account_id):
    """投手诊断：指定账户的项目/广告级诊断"""
    days = int(request.args.get("days", 7))
    return jsonify(radar.promotions(account_id, days))


@app.route("/api/radar/materials/<account_id>")
def api_radar_materials(account_id):
    """投手诊断：指定账户的素材级诊断"""
    days = int(request.args.get("days", 7))
    return jsonify(radar.materials(account_id, days))


@app.route("/api/radar/battle")
def api_radar_battle():
    """每日作战清单：放量/降量/疲劳/陷阱/排查 五张表"""
    days = int(request.args.get("days", 7))
    start_date = request.args.get("start_date")
    end_date = request.args.get("end_date")
    return jsonify(radar.battle_list(days, start_date=start_date, end_date=end_date))


# ── Start ────────────────────────────────────────────────────────

def start_dashboard(port: int = 8888, debug: bool = False):
    logger.info("Starting dual cockpit on http://localhost:%d", port)
    import os as _os
    if not debug or _os.environ.get("WERKZEUG_RUN_MAIN") == "true":
        _start_scheduler()
    app.run(host="0.0.0.0", port=port, debug=debug)

